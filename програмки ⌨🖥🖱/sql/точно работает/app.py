import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkcalendar import DateEntry
from PIL import Image, ImageTk
import sqlite3
import hashlib
from docx import Document
import os
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

class AuthWindow(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Авторизация")
        self.geometry("400x300")
        self.configure_style()

        ttk.Label(self, text="Добро пожаловать!", style="Header.TLabel").pack(pady=20)

        form_frame = ttk.Frame(self)
        form_frame.pack(pady=10)

        ttk.Label(form_frame, text="Email:").grid(row=0, column=0, padx=5, pady=5)
        self.email_entry = ttk.Entry(form_frame, width=25)
        self.email_entry.grid(row=0, column=1)

        ttk.Label(form_frame, text="Пароль:").grid(row=1, column=0, padx=5, pady=5)
        self.password_entry = ttk.Entry(form_frame, show="*", width=25)
        self.password_entry.grid(row=1, column=1)

        ttk.Label(form_frame, text="Роль:").grid(row=2, column=0, padx=5, pady=5)
        self.role_var = tk.StringVar(value="client")
        ttk.Combobox(form_frame, textvariable=self.role_var, 
                    values=["client", "manager", "admin"], state="readonly").grid(row=2, column=1)

        ttk.Button(self, text="Войти", style="Accent.TButton", command=self.login).pack(pady=20)

    def configure_style(self):
        style = ttk.Style()
        style.configure("Header.TLabel", font=("Arial", 16, "bold"))
        style.configure("Accent.TButton", font=("Arial", 12), padding=10)

    def login(self):
        email = self.email_entry.get()
        password = self.password_entry.get()
        role = self.role_var.get()

        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('''
            SELECT * FROM users 
            WHERE email=? AND role=? AND password_hash=?
        ''', (email, role, hash_password(password)))
        
        user = cursor.fetchone()
        conn.close()

        if user:
            self.destroy()
            MainWindow(role, user[0])  # Передаем user_id
        else:
            messagebox.showerror("Ошибка", "Неверные данные")

class MainWindow(tk.Tk):
    def __init__(self, role, user_id):
        super().__init__()
        self.title("Главное меню")
        self.geometry("1024x768")
        self.role = role
        self.user_id = user_id
        self.configure_style()
        self.load_interface()

    def create_order(self):
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        
        try:
            # Создаем заказ
            cursor.execute('''
                INSERT INTO orders (user_id, status, order_date)
                VALUES (?, 'pending', DATE('now'))
            ''', (self.user_id,))
            order_id = cursor.lastrowid
            
            # Переносим товары из корзины в заказ
            cursor.execute('''
                INSERT INTO order_items (order_id, product_id, quantity)
                SELECT ?, product_id, quantity 
                FROM cart 
                WHERE user_id = ?
            ''', (order_id, self.user_id))
            
            # Очищаем корзину
            cursor.execute('DELETE FROM cart WHERE user_id = ?', (self.user_id,))
            
            conn.commit()
            messagebox.showinfo("Успех", "Заказ успешно оформлен!")
            self.load_cart_data()
            
        except sqlite3.Error as e:
            conn.rollback()
            messagebox.showerror("Ошибка", f"Ошибка базы данных: {e}")
            
        finally:
            conn.close()

    def configure_style(self):
        style = ttk.Style()
        style.configure("Treeview.Heading", font=("Arial", 10, "bold"))
        style.map("Treeview", background=[('selected', '#0078D7')])

    def load_interface(self):
        if self.role == "client":
            self.client_interface()
        elif self.role == "manager":
            self.manager_interface()
        elif self.role == "admin":
            self.admin_interface()

    # Клиентский интерфейс
    def client_interface(self):
        notebook = ttk.Notebook(self)
        
        # Вкладка каталога
        catalog_frame = ttk.Frame(notebook)
        self.build_catalog(catalog_frame)
        notebook.add(catalog_frame, text="Каталог")
        
        # Вкладка корзины
        cart_frame = ttk.Frame(notebook)
        self.build_cart(cart_frame)
        notebook.add(cart_frame, text="Корзина")
        
        notebook.pack(expand=True, fill="both")

    def build_catalog(self, parent):
        filter_frame = ttk.Frame(parent)
        filter_frame.pack(pady=10)

        ttk.Label(filter_frame, text="Категория:").pack(side="left")
        self.category_var = tk.StringVar()
        categories = self.get_categories()
        ttk.Combobox(filter_frame, textvariable=self.category_var, 
                    values=categories, state="readonly").pack(side="left", padx=5)

        ttk.Label(filter_frame, text="Цена до:").pack(side="left")
        self.price_filter = ttk.Entry(filter_frame, width=10)
        self.price_filter.pack(side="left", padx=5)
        ttk.Button(filter_frame, text="Применить", command=self.load_products).pack(side="left")

        self.products_frame = ttk.Frame(parent)
        self.products_frame.pack(fill="both", expand=True)
        self.load_products()

    def get_categories(self):
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('SELECT name FROM categories')
        categories = [row[0] for row in cursor.fetchall()]
        conn.close()
        return categories

    def load_products(self):
        for widget in self.products_frame.winfo_children():
            widget.destroy()

        category = self.category_var.get()
        max_price = self.price_filter.get() or 5000

        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('''
            SELECT p.product_id, p.name, p.price, p.image_path 
            FROM products p 
            JOIN categories c ON p.category_id = c.category_id 
            WHERE c.name=? AND p.price <= ?
        ''', (category, float(max_price)))

        row, col = 0, 0
        for product in cursor.fetchall():
            frame = ttk.Frame(self.products_frame, relief="groove", borderwidth=2)
            frame.grid(row=row, column=col, padx=10, pady=10)

            if product[3] and os.path.exists(product[3]):
                try:
                    img = Image.open(product[3])
                    img = img.resize((150, 150), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    img_label = ttk.Label(frame, image=photo)
                    img_label.image = photo
                    img_label.pack()
                except Exception as e:
                    print(f"Ошибка загрузки изображения: {e}")

            ttk.Label(frame, text=product[1], font=("Arial", 10, "bold")).pack()
            ttk.Label(frame, text=f"Цена: {product[2]} руб.").pack()
            ttk.Button(frame, text="В корзину", command=lambda p=product[0]: self.add_to_cart(p)).pack()

            col += 1
            if col > 3:
                col = 0
                row += 1

        conn.close()

    def build_cart(self, parent):
        ttk.Label(parent, text="Корзина", font=("Arial", 14)).pack(pady=10)
        
        columns = ("Товар", "Цена", "Количество")
        self.cart_tree = ttk.Treeview(parent, columns=columns, show="headings")
        for col in columns:
            self.cart_tree.heading(col, text=col)
        self.cart_tree.pack(fill="both", expand=True)
        
        ttk.Button(parent, text="Оформить заказ", command=self.create_order).pack(pady=10)
        self.load_cart_data()

    def load_cart_data(self):
        for row in self.cart_tree.get_children():
            self.cart_tree.delete(row)
        
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('''
            SELECT p.name, p.price, c.quantity 
            FROM cart c 
            JOIN products p ON c.product_id = p.product_id
            WHERE c.user_id = ?
        ''', (self.user_id,))
        
        for item in cursor.fetchall():
            self.cart_tree.insert("", "end", values=item)
        
        conn.close()

    def add_to_cart(self, product_id):
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO cart (user_id, product_id) 
            VALUES (?, ?)
        ''', (self.user_id, product_id))
        conn.commit()
        conn.close()
        self.load_cart_data()
        messagebox.showinfo("Успех", "Товар добавлен в корзину")

    # Интерфейс менеджера
    def manager_interface(self):
        self.title("Панель менеджера")
        
        # Таблица товаров
        columns = ("ID", "Название", "Цена", "Категория", "Остаток")
        self.tree = ttk.Treeview(self, columns=columns, show="headings")
        for col in columns:
            self.tree.heading(col, text=col)
        self.tree.pack(fill="both", expand=True)

        # Кнопки управления
        btn_frame = ttk.Frame(self)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Добавить товар", command=self.add_product).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="Удалить товар", command=self.delete_product).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="Обновить", command=self.load_manager_data).pack(side="left", padx=5)

        self.load_manager_data()

    def load_manager_data(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('''
            SELECT p.product_id, p.name, p.price, c.name, p.stock 
            FROM products p 
            LEFT JOIN categories c ON p.category_id = c.category_id
        ''')
        
        for product in cursor.fetchall():
            self.tree.insert("", "end", values=product)
        
        conn.close()

    def add_product(self):
        add_window = tk.Toplevel(self)
        add_window.title("Добавить товар")

        ttk.Label(add_window, text="Название:").grid(row=0, column=0, padx=5, pady=5)
        name_entry = ttk.Entry(add_window)
        name_entry.grid(row=0, column=1)

        ttk.Label(add_window, text="Цена:").grid(row=1, column=0, padx=5, pady=5)
        price_entry = ttk.Entry(add_window)
        price_entry.grid(row=1, column=1)

        ttk.Label(add_window, text="Категория:").grid(row=2, column=0, padx=5, pady=5)
        category_var = tk.StringVar()
        categories = self.get_categories()
        ttk.Combobox(add_window, textvariable=category_var, values=categories).grid(row=2, column=1)

        ttk.Label(add_window, text="Остаток:").grid(row=3, column=0, padx=5, pady=5)
        stock_entry = ttk.Entry(add_window)
        stock_entry.grid(row=3, column=1)

        ttk.Label(add_window, text="Изображение:").grid(row=4, column=0, padx=5, pady=5)
        image_path = tk.StringVar()
        ttk.Button(add_window, text="Выбрать файл", 
                 command=lambda: image_path.set(filedialog.askopenfilename())).grid(row=4, column=1)

        ttk.Button(add_window, text="Сохранить", 
                 command=lambda: self.save_product(
                     name_entry.get(),
                     price_entry.get(),
                     category_var.get(),
                     stock_entry.get(),
                     image_path.get()
                 )).grid(row=5, columnspan=2, pady=10)

    def save_product(self, name, price, category, stock, image_path):
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        
        cursor.execute('SELECT category_id FROM categories WHERE name=?', (category,))
        category_id = cursor.fetchone()[0]
        
        cursor.execute('''
            INSERT INTO products (name, price, category_id, stock, image_path)
            VALUES (?, ?, ?, ?, ?)
        ''', (name, float(price), category_id, int(stock), image_path))
        
        conn.commit()
        conn.close()
        self.load_manager_data()
        messagebox.showinfo("Успех", "Товар добавлен")

    # Интерфейс администратора
    def admin_interface(self):
        self.title("Панель администратора")
        notebook = ttk.Notebook(self)

        users_tab = ttk.Frame(notebook)
        self.build_users_tab(users_tab)
        notebook.add(users_tab, text="Пользователи")

        reports_tab = ttk.Frame(notebook)
        self.build_reports_tab(reports_tab)
        notebook.add(reports_tab, text="Отчеты")

        notebook.pack(expand=True, fill="both")

    def build_users_tab(self, parent):
        columns = ("ID", "Email", "Роль", "Статус")
        self.users_tree = ttk.Treeview(parent, columns=columns, show="headings")
        for col in columns:
            self.users_tree.heading(col, text=col)
        self.users_tree.pack(fill="both", expand=True)

        btn_frame = ttk.Frame(parent)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Обновить", command=self.load_users).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="Блокировать", command=self.block_user).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="Удалить", command=self.delete_user).pack(side="left", padx=5)

        self.load_users()

    def load_users(self):
        for row in self.users_tree.get_children():
            self.users_tree.delete(row)
        
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('SELECT user_id, email, role, blocked FROM users')
        
        for user in cursor.fetchall():
            status = "Заблокирован" if user[3] else "Активен"
            self.users_tree.insert("", "end", values=(user[0], user[1], user[2], status))
        
        conn.close()

    def block_user(self):
        selected = self.users_tree.selection()
        if not selected:
            return
        
        user_id = self.users_tree.item(selected[0])['values'][0]
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('UPDATE users SET blocked = NOT blocked WHERE user_id=?', (user_id,))
        conn.commit()
        conn.close()
        self.load_users()

    def delete_user(self):
        selected = self.users_tree.selection()
        if not selected:
            return
        
        user_id = self.users_tree.item(selected[0])['values'][0]
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('DELETE FROM users WHERE user_id=?', (user_id,))
        conn.commit()
        conn.close()
        self.load_users()

    def build_reports_tab(self, parent):
        ttk.Label(parent, text="Генерация отчетов", font=("Arial", 14)).pack(pady=10)
        
        date_frame = ttk.Frame(parent)
        date_frame.pack(pady=10)
        
        ttk.Label(date_frame, text="Начало:").grid(row=0, column=0, padx=5)
        self.start_date = DateEntry(date_frame, date_pattern="yyyy-mm-dd")
        self.start_date.grid(row=0, column=1, padx=5)
        
        ttk.Label(date_frame, text="Конец:").grid(row=0, column=2, padx=5)
        self.end_date = DateEntry(date_frame, date_pattern="yyyy-mm-dd")
        self.end_date.grid(row=0, column=3, padx=5)
        
        ttk.Button(parent, text="Экспорт в DOCX", command=self.export_report).pack(pady=10)  # Исправлено

    def export_report(self):
        start_date = self.start_date.get()
        end_date = self.end_date.get()
        
        conn = sqlite3.connect('shop.db')
        cursor = conn.cursor()
        cursor.execute('''
            SELECT o.order_id, u.email, o.status, o.order_date 
            FROM orders o 
            JOIN users u ON o.user_id = u.user_id 
            WHERE o.order_date BETWEEN ? AND ?
        ''', (start_date, end_date))
    
        
        orders = cursor.fetchall()
        conn.close()
        
        doc = Document()
        doc.add_heading('Отчет по заказам', 0)
        doc.add_paragraph(f"Период: {start_date} - {end_date}")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID заказа'
        hdr_cells[1].text = 'Пользователь'
        hdr_cells[2].text = 'Статус'
        hdr_cells[3].text = 'Дата'
        
        for order in orders:
            row_cells = table.add_row().cells
            row_cells[0].text = str(order[0])
            row_cells[1].text = order[1]
            row_cells[2].text = order[2]
            row_cells[3].text = order[3]
        
        doc.save('orders_report.docx')
        messagebox.showinfo("Успех", "Отчет сохранен как orders_report.docx")

if __name__ == "__main__":
    app = AuthWindow()
    app.mainloop()